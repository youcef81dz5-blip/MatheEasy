"""Export extracted Markdown to .md / .tex / .docx.

DOCX export is pure Python (python-docx) — no pandoc required — and converts
LaTeX formulas into native editable Word equations (OMML) via
LaTeX -> MathML -> OMML.
"""

import re
import shutil
import subprocess
import tempfile
import time
from html.parser import HTMLParser
from pathlib import Path

import markdown as md_lib

TEX_TEMPLATE = """% DocExtract export (install pandoc for full LaTeX conversion)
\\documentclass{article}
\\usepackage[utf8]{inputenc}
\\begin{document}

__CONTENT__

\\end{document}
"""

MATH_TOKEN_RE = re.compile("\u27e6M(\\d+)\u27e7")
MD_MATH_RE = re.compile(
    r"\$\$(.+?)\$\$|\\\[(.+?)\\\]|\\\((.+?)\\\)|(?<!\$)\$(?!\$)([^$\n]+?)\$(?!\$)",
    re.S,
)
ARABIC_RE = re.compile(r"[\u0600-\u06FF]")
OMML_NS = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'


def _extract_math(md: str):
    """Replace LaTeX formulas with tokens; return (md, {token_id: latex})."""
    formulas = {}

    def _sub(m):
        latex = next(g for g in m.groups() if g is not None).strip()
        idx = len(formulas)
        formulas[str(idx)] = latex
        return f"\u27e6M{idx}\u27e7"

    return MD_MATH_RE.sub(_sub, md), formulas


def _latex_to_omml(latex: str):
    """LaTeX -> MathML -> OMML xml string, or None on failure."""
    try:
        import latex2mathml.converter
        import mathml2omml

        mathml = latex2mathml.converter.convert(latex)
        omml = mathml2omml.convert(mathml)
        if "<m:oMath" not in omml:
            omml = f"<m:oMath>{omml}</m:oMath>"
        return omml.replace("<m:oMath>", OMML_NS, 1)
    except Exception:
        return None


class _DocxBuilder(HTMLParser):
    """Walk the HTML produced from Markdown and build a Word document."""

    def __init__(self, doc, formulas):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.formulas = formulas
        self.p = None            # current paragraph being filled
        self.bold = False
        self.italic = False
        self.mono = False
        self.list_stack = []     # ['ul'|'ol', ...]
        self.blockquote = False
        self.pre_buffer = None   # collecting <pre> text
        self.rows = None         # buffered table rows
        self.row = None
        self.cell_buf = None

    # -- paragraph helpers -------------------------------------------------
    def _flush(self):
        self.p = None

    def _new_para(self, style=None, center=False):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self._flush()
        self.p = self.doc.add_paragraph(style=style)
        if center:
            self.p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return self.p

    def _add_text_runs(self, text: str):
        if not text or self.p is None:
            return
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        parts = MATH_TOKEN_RE.split(text)
        for i, part in enumerate(parts):
            if i % 2 == 1:                      # math token -> OMML equation
                self._add_math_run(self.formulas.get(part, ""))
                continue
            if not part:
                continue
            run = self.p.add_run(part)
            run.bold = self.bold or None
            run.italic = self.italic or None
            if self.mono:
                run.font.name = "Consolas"
            if ARABIC_RE.search(part) and self.p.alignment is None:
                run.font.rtl = True
                self.p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_math_run(self, latex: str):
        if not latex or self.p is None:
            return
        omml = _latex_to_omml(latex)
        if omml:
            from docx.oxml import parse_xml

            try:
                self.p._p.append(parse_xml(omml))
                return
            except Exception:
                pass
        run = self.p.add_run(latex)             # fallback: raw LaTeX text
        run.italic = True

    # -- tables (buffered so column count is known) -------------------------
    def _start_table(self):
        self.rows = []

    def _build_table(self):
        if not self.rows:
            self.rows = None
            return
        ncols = max(len(r["cells"]) for r in self.rows)
        table = self.doc.add_table(rows=len(self.rows), cols=ncols)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for i, r in enumerate(self.rows):
            for j, content in enumerate(r["cells"]):
                para = table.cell(i, j).paragraphs[0]
                self.p = para
                self.bold = r["header"] or None
                self._add_text_runs(content)
                self.bold = False
        self._flush()
        self.rows = None

    # -- parser handlers -----------------------------------------------------
    def handle_starttag(self, tag, attrs):
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            self._new_para(style=f"Heading {min(level, 4)}")
        elif tag == "p":
            self._new_para(style="Intense Quote" if self.blockquote else None)
        elif tag in ("strong", "b"):
            self.bold = True
        elif tag in ("em", "i"):
            self.italic = True
        elif tag == "code" and self.pre_buffer is None:
            self.mono = True
        elif tag == "br":
            if self.p is not None:
                self.p.add_run().add_break()
        elif tag == "hr":
            self._new_para(center=True).add_run("\u2500" * 40)
            self._flush()
        elif tag in ("ul", "ol"):
            self.list_stack.append(tag)
        elif tag == "li":
            depth = min(len(self.list_stack), 3)
            kind = self.list_stack[-1] if self.list_stack else "ul"
            base = "List Bullet" if kind == "ul" else "List Number"
            style = base if depth == 1 else f"{base} {depth}"
            self._new_para(style=style)
        elif tag == "blockquote":
            self.blockquote = True
        elif tag == "pre":
            self.pre_buffer = []
        elif tag == "table":
            self._start_table()
        elif tag == "tr" and self.rows is not None:
            self.row = {"cells": [], "header": False}
        elif tag in ("td", "th") and self.row is not None:
            if tag == "th":
                self.row["header"] = True
                self.bold = True
            self.cell_buf = []
        elif tag == "img":
            alt = (attrs.get("alt") or "image").strip()
            if self.p is None:
                self._new_para()
            self.p.add_run(f"[🖼 {alt}]")

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._flush()
        elif tag == "p":
            self._flush()
        elif tag in ("strong", "b"):
            self.bold = False
        elif tag in ("em", "i"):
            self.italic = False
        elif tag == "code" and self.pre_buffer is None:
            self.mono = False
        elif tag in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()
        elif tag == "blockquote":
            self.blockquote = False
        elif tag == "pre":
            text = "".join(self.pre_buffer or []).strip("\n")
            self.pre_buffer = None
            p = self._new_para()
            run = p.add_run(text)
            run.font.name = "Consolas"
            self._flush()
        elif tag in ("td", "th"):
            if self.row is not None and self.cell_buf is not None:
                self.row["cells"].append("".join(self.cell_buf).strip())
            self.cell_buf = None
            self.bold = False
        elif tag == "tr" and self.row is not None:
            if self.rows is not None:
                self.rows.append(self.row)
            self.row = None
        elif tag == "table":
            self._build_table()

    def handle_data(self, data):
        if self.pre_buffer is not None:
            self.pre_buffer.append(data)
            return
        if self.cell_buf is not None:
            self.cell_buf.append(data)
            return
        if self.p is None and data.strip():
            self._new_para()
        self._add_text_runs(data)


def _export_docx(md: str, out_path: Path):
    import docx as docx_lib

    doc = docx_lib.Document()
    body_md, formulas = _extract_math(md)
    html = md_lib.markdown(body_md, extensions=["tables", "fenced_code"])
    builder = _DocxBuilder(doc, formulas)
    builder.feed(html)
    builder._flush()
    doc.save(str(out_path))


# ---------------------------------------------------------------- public API

def _pandoc():
    return shutil.which("pandoc")


def _out_dir() -> Path:
    d = Path(tempfile.gettempdir()) / "docextract_exports"
    d.mkdir(parents=True, exist_ok=True)
    return d


def export(md: str, fmt: str) -> str:
    """Export markdown to the requested format, returns the file path."""
    if not md or not md.strip():
        raise ValueError("لا يوجد محتوى للتصدير")

    stamp = time.strftime("%Y%m%d_%H%M%S")
    base = f"docextract_{stamp}"

    if fmt == "md":
        path = _out_dir() / f"{base}.md"
        path.write_text(md, encoding="utf-8")
        return str(path)

    if fmt == "docx":
        path = _out_dir() / f"{base}.docx"
        _export_docx(md, path)
        return str(path)

    if fmt == "tex":
        src = _out_dir() / f"{base}.md"
        src.write_text(md, encoding="utf-8")
        path = _out_dir() / f"{base}.tex"
        pandoc = _pandoc()
        if pandoc:
            subprocess.run(
                [pandoc, str(src), "-f", "markdown", "-t", "latex",
                 "-o", str(path), "--standalone"],
                check=True,
                capture_output=True,
            )
        else:
            path.write_text(TEX_TEMPLATE.replace("__CONTENT__", md), encoding="utf-8")
        return str(path)

    raise ValueError(f"صيغة غير مدعومة: {fmt}")
